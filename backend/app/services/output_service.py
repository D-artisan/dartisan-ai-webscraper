"""
Output service for generating files in various formats.
"""

import json
import os
import uuid
from datetime import datetime
from typing import Dict, Any, Tuple
from pathlib import Path

from docx import Document
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment

from app.utils.logger import logger


class OutputService:
    """Service for generating output files in various formats."""
    
    def __init__(self):
        self.output_dir = Path("outputs")
        self.output_dir.mkdir(exist_ok=True)
    
    async def generate_output(
        self, 
        data: Dict[str, Any], 
        format_type: str, 
        prompt: str
    ) -> Tuple[str, str]:
        """
        Generate output file in the specified format.
        
        Args:
            data: Processed data from LLM
            format_type: Output format (word, pdf, excel, text)
            prompt: Original user prompt
            
        Returns:
            Tuple of (filename, file_path)
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        
        try:
            if format_type == "word":
                return await self._generate_word(data, prompt, timestamp, unique_id)
            elif format_type == "pdf":
                return await self._generate_pdf(data, prompt, timestamp, unique_id)
            elif format_type == "excel":
                return await self._generate_excel(data, prompt, timestamp, unique_id)
            elif format_type == "text":
                return await self._generate_text(data, prompt, timestamp, unique_id)
            else:
                raise ValueError(f"Unsupported format: {format_type}")
                
        except Exception as e:
            logger.error(f"Error generating {format_type} output: {str(e)}")
            raise
    
    async def _generate_word(
        self, 
        data: Dict[str, Any], 
        prompt: str, 
        timestamp: str, 
        unique_id: str
    ) -> Tuple[str, str]:
        """Generate Word document."""
        filename = f"scrape_result_{timestamp}_{unique_id}.docx"
        filepath = self.output_dir / filename
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Web Scraping Results', 0)
        title.alignment = 1  # Center alignment
        
        # Metadata
        doc.add_heading('Scraping Details', level=1)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Prompt: {prompt}")
        
        # Data section
        doc.add_heading('Extracted Data', level=1)
        
        if isinstance(data, dict):
            self._add_dict_to_word(doc, data)
        else:
            doc.add_paragraph(str(data))
        
        doc.save(str(filepath))
        logger.info(f"Generated Word document: {filename}")
        return filename, str(filepath)
    
    def _add_dict_to_word(self, doc: Document, data: Dict[str, Any], level: int = 0):
        """Recursively add dictionary data to Word document."""
        for key, value in data.items():
            if isinstance(value, dict):
                if level == 0:
                    doc.add_heading(str(key).title(), level=2)
                else:
                    doc.add_paragraph(f"{'  ' * level}• {key}:", style='List Bullet')
                self._add_dict_to_word(doc, value, level + 1)
            elif isinstance(value, list):
                doc.add_paragraph(f"{'  ' * level}• {key}:", style='List Bullet')
                for item in value:
                    if isinstance(item, dict):
                        self._add_dict_to_word(doc, item, level + 1)
                    else:
                        doc.add_paragraph(f"{'  ' * (level + 1)}- {item}")
            else:
                doc.add_paragraph(f"{'  ' * level}• {key}: {value}")
    
    async def _generate_pdf(
        self, 
        data: Dict[str, Any], 
        prompt: str, 
        timestamp: str, 
        unique_id: str
    ) -> Tuple[str, str]:
        """Generate PDF document."""
        filename = f"scrape_result_{timestamp}_{unique_id}.pdf"
        filepath = self.output_dir / filename
        
        doc = SimpleDocTemplate(str(filepath), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            spaceAfter=30,
            alignment=1  # Center
        )
        story.append(Paragraph("Web Scraping Results", title_style))
        story.append(Spacer(1, 20))
        
        # Metadata
        story.append(Paragraph("Scraping Details", styles['Heading2']))
        story.append(Paragraph(f"<b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Paragraph(f"<b>Prompt:</b> {prompt}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Data
        story.append(Paragraph("Extracted Data", styles['Heading2']))
        
        if isinstance(data, dict) and self._is_tabular_data(data):
            # Create table for tabular data
            table_data = self._dict_to_table_data(data)
            if table_data:
                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
        else:
            # Add as formatted text
            formatted_text = self._format_data_for_pdf(data)
            story.append(Paragraph(formatted_text, styles['Normal']))
        
        doc.build(story)
        logger.info(f"Generated PDF document: {filename}")
        return filename, str(filepath)
    
    def _is_tabular_data(self, data: Dict[str, Any]) -> bool:
        """Check if data can be represented as a table."""
        if not isinstance(data, dict):
            return False
        
        # Look for list of dictionaries (common table format)
        for value in data.values():
            if isinstance(value, list) and len(value) > 0:
                if all(isinstance(item, dict) for item in value):
                    return True
        return False
    
    def _dict_to_table_data(self, data: Dict[str, Any]) -> list:
        """Convert dictionary data to table format."""
        table_data = []
        
        # Find the first list of dictionaries
        for key, value in data.items():
            if isinstance(value, list) and len(value) > 0:
                if all(isinstance(item, dict) for item in value):
                    # Use this as table data
                    if value:
                        headers = list(value[0].keys())
                        table_data.append(headers)
                        for item in value:
                            row = [str(item.get(header, '')) for header in headers]
                            table_data.append(row)
                    break
        
        return table_data
    
    def _format_data_for_pdf(self, data: Any, indent: int = 0) -> str:
        """Format data for PDF display."""
        if isinstance(data, dict):
            result = []
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    result.append(f"{'&nbsp;' * (indent * 4)}<b>{key}:</b><br/>")
                    result.append(self._format_data_for_pdf(value, indent + 1))
                else:
                    result.append(f"{'&nbsp;' * (indent * 4)}<b>{key}:</b> {value}<br/>")
            return ''.join(result)
        elif isinstance(data, list):
            result = []
            for item in data:
                result.append(f"{'&nbsp;' * (indent * 4)}• {self._format_data_for_pdf(item, indent + 1)}<br/>")
            return ''.join(result)
        else:
            return str(data)
    
    async def _generate_excel(
        self, 
        data: Dict[str, Any], 
        prompt: str, 
        timestamp: str, 
        unique_id: str
    ) -> Tuple[str, str]:
        """Generate Excel file."""
        filename = f"scrape_result_{timestamp}_{unique_id}.xlsx"
        filepath = self.output_dir / filename
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Scraping Results"
        
        # Header styling
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        # Add metadata
        ws['A1'] = "Web Scraping Results"
        ws['A1'].font = Font(bold=True, size=16)
        ws['A3'] = "Generated:"
        ws['B3'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        ws['A4'] = "Prompt:"
        ws['B4'] = prompt
        
        # Add data starting from row 6
        current_row = 6
        
        if isinstance(data, dict):
            current_row = self._add_dict_to_excel(ws, data, current_row)
        else:
            ws[f'A{current_row}'] = "Extracted Data"
            ws[f'A{current_row}'].font = header_font
            ws[f'A{current_row}'].fill = header_fill
            current_row += 1
            ws[f'A{current_row}'] = str(data)
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(str(filepath))
        logger.info(f"Generated Excel file: {filename}")
        return filename, str(filepath)
    
    def _add_dict_to_excel(self, ws, data: Dict[str, Any], start_row: int) -> int:
        """Add dictionary data to Excel worksheet."""
        current_row = start_row
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        for key, value in data.items():
            if isinstance(value, list) and len(value) > 0:
                if all(isinstance(item, dict) for item in value):
                    # Create table for list of dictionaries
                    ws[f'A{current_row}'] = key.title()
                    ws[f'A{current_row}'].font = header_font
                    ws[f'A{current_row}'].fill = header_fill
                    current_row += 1
                    
                    if value:
                        headers = list(value[0].keys())
                        # Add headers
                        for col, header in enumerate(headers, 1):
                            cell = ws.cell(row=current_row, column=col, value=header)
                            cell.font = Font(bold=True)
                            cell.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
                        current_row += 1
                        
                        # Add data rows
                        for item in value:
                            for col, header in enumerate(headers, 1):
                                ws.cell(row=current_row, column=col, value=str(item.get(header, '')))
                            current_row += 1
                        current_row += 1  # Add space between sections
                else:
                    # Simple list
                    ws[f'A{current_row}'] = key.title()
                    ws[f'A{current_row}'].font = header_font
                    ws[f'A{current_row}'].fill = header_fill
                    current_row += 1
                    for item in value:
                        ws[f'A{current_row}'] = str(item)
                        current_row += 1
                    current_row += 1
            else:
                # Simple key-value pair
                ws[f'A{current_row}'] = key.title()
                ws[f'B{current_row}'] = str(value)
                current_row += 1
        
        return current_row
    
    async def _generate_text(
        self, 
        data: Dict[str, Any], 
        prompt: str, 
        timestamp: str, 
        unique_id: str
    ) -> Tuple[str, str]:
        """Generate text file."""
        filename = f"scrape_result_{timestamp}_{unique_id}.txt"
        filepath = self.output_dir / filename
        
        content = []
        content.append("=" * 50)
        content.append("WEB SCRAPING RESULTS")
        content.append("=" * 50)
        content.append("")
        content.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        content.append(f"Prompt: {prompt}")
        content.append("")
        content.append("EXTRACTED DATA:")
        content.append("-" * 20)
        content.append("")
        
        if isinstance(data, dict):
            content.extend(self._format_dict_for_text(data))
        else:
            content.append(str(data))
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
        
        logger.info(f"Generated text file: {filename}")
        return filename, str(filepath)
    
    def _format_dict_for_text(self, data: Dict[str, Any], indent: int = 0) -> list:
        """Format dictionary data for text output."""
        lines = []
        indent_str = "  " * indent
        
        for key, value in data.items():
            if isinstance(value, dict):
                lines.append(f"{indent_str}{key.upper()}:")
                lines.extend(self._format_dict_for_text(value, indent + 1))
                lines.append("")
            elif isinstance(value, list):
                lines.append(f"{indent_str}{key.upper()}:")
                for i, item in enumerate(value, 1):
                    if isinstance(item, dict):
                        lines.append(f"{indent_str}  {i}.")
                        lines.extend(self._format_dict_for_text(item, indent + 2))
                    else:
                        lines.append(f"{indent_str}  {i}. {item}")
                lines.append("")
            else:
                lines.append(f"{indent_str}{key}: {value}")
        
        return lines
    
    def cleanup_old_files(self, days: int = 7):
        """Clean up old output files."""
        try:
            import time
            current_time = time.time()
            for file_path in self.output_dir.glob("*"):
                if file_path.is_file():
                    file_age = current_time - file_path.stat().st_mtime
                    if file_age > (days * 24 * 3600):  # Convert days to seconds
                        file_path.unlink()
                        logger.info(f"Cleaned up old file: {file_path.name}")
        except Exception as e:
            logger.error(f"Error cleaning up files: {str(e)}")
